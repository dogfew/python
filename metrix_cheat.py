import os
import subprocess
import pandas as pd
import re
import numpy as np
from docx import Document
from statsmodels.api import OLS
from scipy.stats import ttest_1samp
from sklearn.linear_model import LinearRegression
from tqdm import tqdm


def get_res(first_task, second_task,
            third_task, fourth_task, five_task,
            sixth_task, seventh_task,
            final_task,
            data="myData.csv"):
    df = pd.read_csv(data).assign(X0=1)
    reg = OLS(df.Y, df[['X0', 'X1', 'X2']]).fit()
    answers = [
        df[first_task].min(),

        df[second_task].var(),

        df[df[third_task[0]] < float(third_task[2])].Y.mean() if 'меньше' in third_task[1] else
        df[df[third_task[0]] > float(third_task[2])].Y.mean(),

        ttest_1samp(df.Y, float(fourth_task)).pvalue,

        df[five_task[0]].corr(df[five_task[1]]),

        reg.params[
            2 if sixth_task == 'X2' else (1 if sixth_task == 'X1' else 0)
        ],

        reg.rsquared,
        reg.tvalues[
            2 if seventh_task == 'X2' else (1 if seventh_task == 'X1' else 0)
        ],
        reg.ssr / reg.df_resid,
        LinearRegression(fit_intercept=False).fit(df[['X1', 'X2']], df.Y).coef_[
            1 if final_task == 'X2' else 0]
    ]
    return answers


path = "" # path to docx files 
os.chdir(path)
documents = [Document(x) for x in os.listdir() if x.endswith(".docx")]
total_answers = []
total_names = []
for doc in tqdm(documents):
    p = doc.paragraphs
    name = re.match(r"Привет тебе, (.*?)!", doc.paragraphs[0].text).group(0)[13:-1]
    r_script_data = '\n'.join(list(map(lambda x: x.text, doc.paragraphs[2:7])))
    r_script_data += '\nwrite.csv(myData, "myData.csv", row.names = F)'
    with open('script.r', 'w+') as f:
        f.writelines(r_script_data)
    subprocess.call(['Rscript', 'script.r'])
    first_task = re.findall('значение (.*?) в', doc.paragraphs[18].text)[0]
    second_task = re.findall('дисперсию (.*?) в', doc.paragraphs[21].text)[0]
    third_task = re.findall(', где (.*)', doc.paragraphs[24].text)[0].split()
    fourth_task = re.findall('равно (.*). В ответ', doc.paragraphs[27].text)[0]
    five_task = re.findall('корреляцию (.*) и (.*)', doc.paragraphs[30].text)[0]
    if "констант" not in doc.paragraphs[33].text:
        sixth_task = re.findall('при (.*) в регрессии', doc.paragraphs[33].text)[0]
    else:
        sixth_task = 'констант'
    if "констант" in doc.paragraphs[39].text:
        seventh_task = 'конст'
    else:
        seventh_task = re.findall('о равенстве коэффициента при (.*) нулю в регрессии', doc.paragraphs[39].text)[0]
    final_task = re.findall('коэффициент при (.*) в регрессии', doc.paragraphs[45].text)[0]
    ans = get_res(first_task, second_task, third_task,
                  fourth_task, five_task,  sixth_task, seventh_task, final_task)
    total_answers.append(ans)
    total_names.append(name)

res = pd.DataFrame(np.array(total_answers).T, columns=total_names).round(3).T
res.columns = np.arange(1, 11, 1)
os.chdir('..')
res.to_excel('answers.xlsx')
print(res)